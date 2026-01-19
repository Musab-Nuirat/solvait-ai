"""Document Ingestion Pipeline - Supports PDF and DOCX with semantic chunking."""

import os
from pathlib import Path
from typing import List, Optional

from llama_index.core import (
    VectorStoreIndex,
    StorageContext,
)
from llama_index.core.node_parser import SentenceSplitter
from llama_index.core.schema import Document
from app.agent.custom_embedding import GoogleGenaiEmbedding
from llama_index.vector_stores.chroma import ChromaVectorStore
import chromadb

from app.config import MANUALS_DIR, CHROMA_DIR, GOOGLE_API_KEY, LLAMA_CLOUD_API_KEY


# Try to import LlamaParse (optional, falls back to simple parsing)
try:
    from llama_parse import LlamaParse
    LLAMA_PARSE_AVAILABLE = True
except ImportError:
    LLAMA_PARSE_AVAILABLE = False
    print("LlamaParse not available. Using fallback parsing.")


class DocumentIngestion:
    """Handle document ingestion for the RAG pipeline."""

    def __init__(self):
        print("Initializing DocumentIngestion...")
        print("Setting up GoogleGenaiEmbedding...")
        self.embed_model = GoogleGenaiEmbedding(
            model_name="models/embedding-001",
            api_key=GOOGLE_API_KEY
        )

        # Use SentenceSplitter for better semantic chunking of HR manuals
        # - chunk_size: 1024 characters (good for policy sections)
        # - chunk_overlap: 200 characters (preserves context between chunks)
        print("Setting up SentenceSplitter for semantic chunking...")
        self.node_parser = SentenceSplitter(
            chunk_size=1024,
            chunk_overlap=200,
            paragraph_separator="\n\n",
        )

        # Initialize ChromaDB
        print(f"Connecting to ChromaDB at {CHROMA_DIR}...")
        self.chroma_client = chromadb.PersistentClient(path=str(CHROMA_DIR))
        print("Getting/creating collection 'hr_policies'...")
        self.collection = self.chroma_client.get_or_create_collection("hr_policies")
        print("Wrapping collection with ChromaVectorStore...")
        self.vector_store = ChromaVectorStore(chroma_collection=self.collection)
        print("DocumentIngestion initialized.")

    def parse_docx(self, docx_path: Path) -> str:
        """
        Parse DOCX file to plain text with structure preservation.
        """
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(str(docx_path))
            content_parts = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                # Detect headings by style
                style_name = para.style.name.lower() if para.style else ""

                if "heading 1" in style_name or "title" in style_name:
                    content_parts.append(f"\n# {text}\n")
                elif "heading 2" in style_name:
                    content_parts.append(f"\n## {text}\n")
                elif "heading 3" in style_name:
                    content_parts.append(f"\n### {text}\n")
                elif "heading" in style_name:
                    content_parts.append(f"\n#### {text}\n")
                else:
                    content_parts.append(text)

            # Also extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        table_text.append(row_text)
                if table_text:
                    content_parts.append("\n" + "\n".join(table_text) + "\n")

            return "\n".join(content_parts)

        except Exception as e:
            print(f"Error parsing DOCX {docx_path}: {e}")
            return f"# Document: {docx_path.name}\n\nError parsing document: {str(e)}"

    def parse_pdf_to_markdown(self, pdf_path: Path) -> str:
        """
        Parse PDF to Markdown using LlamaParse.
        Falls back to simple extraction if LlamaParse unavailable.
        """
        if LLAMA_PARSE_AVAILABLE and LLAMA_CLOUD_API_KEY:
            parser = LlamaParse(
                api_key=LLAMA_CLOUD_API_KEY,
                result_type="markdown",
                language="en",
                verbose=True
            )
            documents = parser.load_data(str(pdf_path))
            return "\n\n".join([doc.text for doc in documents])
        else:
            try:
                import fitz  # PyMuPDF
                doc = fitz.open(str(pdf_path))
                text = ""
                for page in doc:
                    text += page.get_text("text") + "\n\n"
                return text
            except ImportError:
                return f"# Document: {pdf_path.name}\n\nPlease install PyMuPDF or LlamaParse for PDF parsing."

    def ingest_manuals(self, force_reindex: bool = False) -> VectorStoreIndex:
        """
        Ingest all PDF and DOCX manuals from the manuals directory.

        Args:
            force_reindex: If True, reindex even if index exists

        Returns:
            VectorStoreIndex ready for querying
        """
        # Check if index already exists
        index_exists = len(self.collection.get()["ids"]) > 0

        if index_exists and not force_reindex:
            print("Using existing index from ChromaDB")
            return VectorStoreIndex.from_vector_store(
                self.vector_store,
                embed_model=self.embed_model
            )

        # Find all supported files in manuals directory
        pdf_files = list(MANUALS_DIR.glob("*.pdf"))
        docx_files = list(MANUALS_DIR.glob("*.docx"))
        all_files = pdf_files + docx_files

        if not all_files:
            print(f"No PDF or DOCX files found in {MANUALS_DIR}")
            # Create empty index with sample policy
            return self._create_index_with_sample()

        print(f"Found {len(all_files)} files to process ({len(pdf_files)} PDF, {len(docx_files)} DOCX)")

        documents = []

        # Always include sample policy as fallback (contains actual policy rules)
        sample_doc = Document(
            text=SAMPLE_POLICY_MARKDOWN,
            metadata={"source": "sample_policy_handbook", "type": "policy_manual", "language": "en"}
        )
        documents.append(sample_doc)
        print("  Added: sample_policy_handbook (fallback HR policies)")

        for file_path in all_files:
            print(f"  Processing: {file_path.name}")

            if file_path.suffix.lower() == ".pdf":
                content = self.parse_pdf_to_markdown(file_path)
            elif file_path.suffix.lower() == ".docx":
                content = self.parse_docx(file_path)
            else:
                continue

            # Detect language from filename
            is_arabic = "arabic" in file_path.name.lower() or "عربي" in file_path.name

            doc = Document(
                text=content,
                metadata={
                    "source": file_path.name,
                    "type": "policy_manual",
                    "language": "ar" if is_arabic else "en"
                }
            )
            documents.append(doc)
            print(f"    -> Extracted {len(content)} characters, language: {'Arabic' if is_arabic else 'English'}")

        return self._create_index(documents)

    def _create_index_with_sample(self) -> VectorStoreIndex:
        """Create index with sample policy when no files exist."""
        doc = Document(
            text=SAMPLE_POLICY_MARKDOWN,
            metadata={"source": "sample_handbook", "type": "policy_manual", "language": "en"}
        )
        return self._create_index([doc], force_clear=True)

    def ingest_markdown_directly(self, markdown_content: str, source_name: str) -> VectorStoreIndex:
        """
        Ingest markdown content directly (for testing without files).
        """
        doc = Document(
            text=markdown_content,
            metadata={
                "source": source_name,
                "type": "policy_manual"
            }
        )
        return self._create_index([doc], force_clear=True)

    def _create_index(self, documents: List[Document], force_clear: bool = False) -> VectorStoreIndex:
        """Create vector index from documents."""
        if force_clear:
            # Clear existing collection
            try:
                self.chroma_client.delete_collection("hr_policies")
            except Exception:
                pass
            self.collection = self.chroma_client.create_collection("hr_policies")
            self.vector_store = ChromaVectorStore(chroma_collection=self.collection)

        # Parse documents into nodes using SentenceSplitter
        nodes = self.node_parser.get_nodes_from_documents(documents)
        print(f"Created {len(nodes)} chunks using SentenceSplitter")

        # Log chunk distribution
        for doc in documents:
            source = doc.metadata.get("source", "unknown")
            doc_nodes = [n for n in nodes if n.metadata.get("source") == source]
            print(f"  -> {source}: {len(doc_nodes)} chunks")

        # Create index
        storage_context = StorageContext.from_defaults(vector_store=self.vector_store)
        index = VectorStoreIndex(
            nodes,
            storage_context=storage_context,
            embed_model=self.embed_model
        )

        print("Index created and persisted to ChromaDB")
        return index


# Sample policy content for testing
SAMPLE_POLICY_MARKDOWN = """
# PeopleHub Employee Handbook

## Chapter 1: Leave Policies

### 1.1 Annual Leave
- All employees are entitled to 21 days of paid annual leave per year.
- New employees (less than 1 year) receive 14 days.
- Leave requests must be submitted at least 3 days in advance.
- Maximum consecutive leave is 10 working days without special approval.

### 1.2 Sick Leave
- Employees receive 10 days of paid sick leave per year.
- Medical certificate required for sick leave exceeding 2 days.
- Unused sick leave does not carry over to the next year.

### 1.3 Unpaid Leave
- Unpaid leave may be granted for personal circumstances.
- Maximum 30 days per year with manager approval.
- Does not affect annual leave balance.

## Chapter 2: Compensation & Benefits

### 2.1 Salary Structure
- Salaries are paid monthly on the last working day.
- Salary breakdown includes: Basic Salary, Housing Allowance, Transport Allowance.

### 2.2 Overtime Policy
- Overtime is paid at 1.5x the hourly rate for weekdays.
- Weekend overtime is paid at 2x the hourly rate.
- Overtime requires prior manager approval.
- This policy applies to non-management roles only.

### 2.3 Salary Advance
- Employees may request a salary advance up to 50% of monthly salary.
- Eligibility: Minimum 1 year of continuous employment required.
- Advance is deducted from the next salary payment.
- Maximum one advance per quarter.

## Chapter 3: Attendance & Punctuality

### 3.1 Working Hours
- Standard working hours: 8:00 AM to 5:00 PM, Sunday to Thursday.
- 1 hour lunch break from 12:00 PM to 1:00 PM.

### 3.2 Late Arrival Policy
- Employees must submit an excuse request for arrivals after 8:15 AM.
- Frequent tardiness (more than 3 times per month) may affect performance reviews.

### 3.3 Early Departure
- Early departure requires advance approval or same-day excuse submission.
- Reason must be provided for all early departure requests.

## Chapter 4: Health Insurance

### 4.1 Coverage
- Medical insurance covers employees and immediate family members.
- Dental coverage is included in the standard plan.
- Vision coverage requires additional premium.

### 4.2 Claims Process
- Submit claims through the HR portal within 30 days of treatment.
- Pre-approval required for procedures exceeding SAR 5,000.
"""


def get_sample_policy_index() -> VectorStoreIndex:
    """Get or create an index with sample policy content for testing."""
    ingestion = DocumentIngestion()
    return ingestion.ingest_markdown_directly(
        SAMPLE_POLICY_MARKDOWN,
        "sample_employee_handbook.md"
    )
